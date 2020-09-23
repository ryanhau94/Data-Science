# Requires python-docx, which can be installed by 'pip install python-docx'

# -*- coding: UTF-8 -*-
 
import requests
import urllib.request
from bs4 import BeautifulSoup
import os
import time
from queue import Queue

from docx import Document
from docx.shared import Inches

if __name__ == '__main__':
    folder_name = 'temperary'
    cwd = os.getcwd()
    #folder_path = os.path.join(cwd, folder_name)
    folder_path = './temperary/'
    if not os.path.exists(folder_path):
        os.mkdir(folder_path)

    document = Document()
    document.add_heading('PAL3 Curio Walkthrough', 0)

    url_queue = Queue()
    url_init = 'https://wap.gamersky.com/gl/Content-785240.html'
    url_seen = set()

    headers = {'User-Agent': 'Mozilla/5.0'}

    url_seen.add(url_init)
    url_queue.put(url_init)

    cnt = 0

    while(True):
        url = url_queue.get()
        print('Start processing site %d ...' %(cnt))
        #print(url)
        try:
            response = requests.get(url, headers=headers)
            soup = BeautifulSoup(response.content, 'html.parser')
            for item in soup.find_all('a'):
                try:
                    if item.get('class')[0] == 'page-next':
                        url_next = item.get('href')
                        break
                except:
                    continue
            
            if url_next not in url_seen:
                url_seen.add(url_next)
                url_queue.put(url_next)
            
            img_html = requests.get(soup.find_all('img')[0].get('src')) 
            img_name = folder_path + 'img_' + str(cnt) +'.png'
            with open(img_name, 'wb') as file:  
                file.write(img_html.content)
                file.flush()
            file.close()

            putting = False
            for item in soup.find_all('p'):
                if putting:
                    document.add_paragraph(list(item)[0])
                if len(item.find_all('strong')) > 0:
                    putting = True
                    document.add_heading(list(item)[0], level=1)
                if str(item).find('如图：') >= 0:
                    break
            document.add_picture(img_name)

        except:
            print('Unkown Error: May not affect the funcationality')
        
        print('Done')
        time.sleep(1)
        cnt = cnt+1
        if url_queue.qsize() <= 0:
            break

    document.save('PAL3_Curio_Walkthrough.docx')
    #os.rmdir(folder_path)

    print('All Done!')